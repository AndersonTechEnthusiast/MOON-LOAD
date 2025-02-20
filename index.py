import pytubefix
from pytubefix.cli import on_progress
from pytubefix import YouTube
import os
import time
import whisper
import reportlab
from reportlab.platypus import SimpleDocTemplate
from reportlab.platypus import Paragraph
from reportlab.platypus import Image
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.enums import TA_RIGHT
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.enums import TA_JUSTIFY
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from datetime import datetime
import validators
import requests
import xml.etree.ElementTree 
txt = "reset.txt"
_KEYSWORDS = []
# VERIFICA SE O DOCUMENTO EXISTE
if os.path.exists(txt):
    # SE EXISTIR ABRA-O , NO MODO `LEITURA` , com enconding="uft-8" 
    with open(txt , "r" , encoding="utf-8") as ler_documento:
        # ler o texto , .read() -> leia
        texto_lido = ler_documento.read()
        # verifica se o _RESET_ esta contido nele
        if "_RESET_" in texto_lido:
            # se estiver 
            # replace = substitua o '_RESET_' por '' vazio , e tira os espaços possiveis existentes .strip()
            number = texto_lido.replace('_RESET_','').strip()
            # aloca no proprio numero ele mesmo como um número int()
            number = int(number)
            # soma mais 1 nop proprio número
            number = number + 1
            # escreva novamente no documento com o number atualizado
            with open(txt , "w" , encoding="utf-8") as documento:
                # escreva
                documento.write(f"_RESET_{number}")
# SE NÃO EXISTIR 
if not os.path.exists(txt):
    # CRIE
    with open(txt , "w" , encoding="utf-8") as documento:
        documento.write(f"_RESET_1")
# MOON LOAD #
Moon_Load_Loading = ['🌑','🌒','🌓','🌔','🌕','🌖','🌗','🌘']
Moon_Load_Loading_Texts = {
    'criar_diretório': {
        'módulo-1':['Trabalhando Nisso','Trabalhando Nisso.','Trabalhando Nisso..','Trabalhando Nisso...'],
        'módulo-2':['Criando Diretórios Para Seus Documentos Words','Criando Diretórios Para Seus Documentos Words.','Criando Diretórios Para Seus Documentos Words..','Criando Diretórios Para Seus Documentos Words...'],
        'módulo-3':['Criando Diretórios Para Seus Documentos PDFs','Criando Diretórios Para Seus Documentos PDFs.','Criando Diretórios Para Seus Documentos PDFs..','Criando Diretórios Para Seus Documentos PDFs...'],
    },
    'processamento_de_link':{
        'módulo-1':['Verificando Link Informado','Verificando Link Informado.','Verificando Link Informado..','Verificando Link Informado...'],
        'módulo-2':['Processando Link Informado','Processando Link Informado.','Processando Link Informado..','Processando Link Informado...'],
        'módulo-3':['Buscando Vídeo do Link Informado','Buscando Vídeo do Link Informado.','Buscando Vídeo do Link Informado..','Buscando Vídeo do Link Informado...'],
    },
    'processamento_de_video': {
        'módulo-1':['Preparando para o Download do Vídeo','Preparando para o Download do Vídeo.','Preparando para o Download do Vídeo..','Preparando para o Download do Vídeo...'],
        'módulo-2':['Fazendo Download do Vídeo Informado','Fazendo Download do Vídeo Informado.','Fazendo Download do Vídeo Informado..','Fazendo Download do Vídeo Informado...'],
        'módulo-3':['Extraindo Todas as Informações do Vídeo','Extraindo Todas as Informações do Vídeo.','Extraindo Todas as Informações do Vídeo..','Extraindo Todas as Informações do Vídeo...'],
        'módulo-4':['Adicionando o Vídeo no seu Diretório Criado','Adicionando o Vídeo no seu Diretório Criado.','Adicionando o Vídeo no seu Diretório Criado..','Adicionando o Vídeo no seu Diretório Criado...'],
    },
    'processamento_de_video_not_audio': {
        'módulo-1':['Preparando para o Download do Vídeo (sem Aúdio)','Preparando para o Download do Vídeo (sem Aúdio).','Preparando para o Download do Vídeo (sem Aúdio)..','Preparando para o Download do Vídeo (sem Aúdio)...'],
        'módulo-2':['Fazendo Download do Vídeo (sem Aúdio) Informado','Fazendo Download do Vídeo (sem Aúdio) Informado.','Fazendo Download do Vídeo (sem Aúdio) Informado..','Fazendo Download do Vídeo (sem Aúdio) Informado...'],
        'módulo-3':['Extraindo Todas as Informações do Vídeo (sem Aúdio)','Extraindo Todas as Informações do Vídeo (sem Aúdio).','Extraindo Todas as Informações do Vídeo (sem Aúdio)..','Extraindo Todas as Informações do Vídeo (sem Aúdio)...'],
        'módulo-4':['Adicionando o Vídeo (sem Aúdio) no seu Diretório Criado','Adicionando o Vídeo (sem Aúdio) no seu Diretório Criado.','Adicionando o Vídeo (sem Aúdio) no seu Diretório Criado..','Adicionando o Vídeo (sem Aúdio) no seu Diretório Criado...'],
    },
    'processamento_de_audio': {
        'módulo-1':['Preparando para o Download do Áudio','Preparando para o Download do Áudio.','Preparando para o Download do Áudio..','Preparando para o Download do Áudio...'],
        'módulo-2':['Fazendo Download do Áudio Informado','Fazendo Download do Áudio Informado.','Fazendo Download do Áudio Informado..','Fazendo Download do Áudio Informado...'],
        'módulo-3':['Extraindo Todas as Informações do Áudio','Extraindo Todas as Informações do Áudio.','Extraindo Todas as Informações do Áudio..','Extraindo Todas as Informações do Áudio...'],
        'módulo-4':['Adicionando o Áudio no seu Diretório Criado','Adicionando o Áudio no seu Diretório Criado.','Adicionando o Áudio no seu Diretório Criado..','Adicionando o Áudio no seu Diretório Criado...'],
    },
    'processamento_da_transcrição': {
        'módulo-1':['Preparando Áudio para a Transcrição','Preparando Áudio para a Transcrição.','Preparando Áudio para a Transcrição..','Preparando Áudio para a Transcrição...'],
        'módulo-2':['Transcrevendo Áudio','Transcrevendo Áudio.','Transcrevendo Áudio..','Transcrevendo Áudio...'],   
    },
    'processamento_da_criação_de_documento_word': {
        'módulo-1':['Criando Documento Word','Criando Documento Word.','Criando Documento Word..','Criando Documento Word...'],
        'módulo-2':['Escrevendo no Documento Word','Escrevendo no Documento Word.','Escrevendo no Documento Word..','Escrevendo no Documento Word...'],
    },
    'processamento_da_criação_de_documento_pdf': {
        'módulo-1':['Criando Documento PDF','Criando Documento PDF.','Criando Documento PDF..','Criando Documento PDF...'],
        'módulo-2':['Escrevendo no Documento PDF','Escrevendo no Documento PDF.','Escrevendo no Documento PDF..','Escrevendo no Documento PDF...'],
    }
}
# datetime = puxa a data e a hora
# .now() = agora , puxa a data e a hora de agora
# hour = puxa apenas a hora , tipo: se for 13:25:45 , ele retornará apenas 13
hora = datetime.now().hour

# verifica se for menor ou igual a 12 , é manhã , 'Bom Dia'
if int(hora) <= 12:
    # imprime a mensagem
    time.sleep(1)
    print(f"👋 Olá , Bom dia !!!")
# verifica se for maior que 12 , é tarde , 'Boa Tarde'
elif int(hora) > 12 and int(hora) <= 17:
    # imprime a mensagem
    time.sleep(1)
    print(f"👋 Olá , Boa tarde !!!")
# se não
else:
    # imprime a mensagem
    time.sleep(1)
    print(f"👋 Olá , Boa noite !!!")

time.sleep(1)

diretorio = input(" 📁  Diretório: ")
while not diretorio:
    diretorio = input(" ❌  Error , Diretório não pode ser Nulo \n  📁 Diretório: ")
while os.path.exists(diretorio):
    diretorio = input(" ❌  Error , Diretório já existe \n  📁 Diretório: ")
diretorio = diretorio.strip()
# contador , é igual a 0
contador = 0
# inicia no modulo 1 , caso não entenda , veja o  Moon_Load_Loading_Texts['criar_diretório']['módulo-1'] , conforme percorrer , incrementa o contador e o modulo
modulo = 1
# enquanto True , ou seja enquanto verdade , seja verdade , faça o looping
while True:
    # zip() , é um interador de conjunto de arrays , me permitindo percorrer duas arrays ao mesmo tempo;
    # zip(array1 , array2)
    # para i , e em zip(Array1 , Array2)
    # i percorre a Array1 = i = variavel (poderia qualquer outra coisa)
    # t percorre a Array2 = t = variavel (poderia qualquer outra coisa)
    for i , t in zip(Moon_Load_Loading , Moon_Load_Loading_Texts['criar_diretório'][f'módulo-{modulo}']):
        # conta 1 milisegundo
        time.sleep(.1)
        # incrimenta o contador a cada loop
        contador = contador + 1
        # imprime a mensagem mostrando o dado de cada array
        print(f"\r {i} {t} | {contador}%             ", end="")
        # conta 1 milisegundo
        time.sleep(.1)
        # enquanto contador for igual a 100
        while contador == 100:
            # reinicia o contador igualando ele a 0 novamente
            contador = 0
            # cria o diretório
            if modulo == 1:
                os.makedirs(diretorio)
            if modulo == 2:
                os.makedirs(f"{diretorio}-Word")
            if modulo == 3:
                os.makedirs(f"{diretorio}-PDF")
            # incrementa o modulo , de 1 para 2 , de 2 para 3 , de 3 para 4 , (os módulos da array 'criar_diretório' , vão apenas até o módulo 3 , mas paramos no 4 pois ele para antes de percorrer , então após o 3 incrementa e para (break) )
            modulo = modulo + 1
            # se for igual a 4 
            if modulo == 4:
                # quebra o looping
                break
    # se o modulo for igual a 4 
    if modulo == 4:
        modulo = 1
        contador = 0
        # quebra o looping 
        break
print("\n")
print("🔚")
break_question_progress = input(f"o 📁  Diretório Criado {diretorio} \n é o Desejado ? \n 🟢  [S] para sim \n 🔴  [N] para não \n Resposta: ")

while not break_question_progress:
    break_question_progress = input(f"❌ Error , Resposta Inválida Certifique-se de que sua Resposta Não esteja Vazia e Corresponda a sim ou não \n 📁  Diretório Criado {diretorio} \n é o Desejado ? \n 🟢  [S] para sim \n 🔴  [N] para não \n Resposta: ")
while break_question_progress.lower() != 'n' and break_question_progress.lower() != 's':
    break_question_progress = input(f"❌ Error , Resposta Inválida Certifique-se de que sua Resposta Não esteja Vazia e Corresponda a sim ou não \n 📁  Diretório Criado {diretorio} \n é o Desejado ? \n 🟢  [S] para sim \n 🔴  [N] para não \n Resposta: ")
while break_question_progress.lower() == 'n':  
    if os.path.exists(diretorio):
        os.removedirs(diretorio)
        os.removedirs(f"{diretorio}-Word")
        os.removedirs(f"{diretorio}-PDF")
    diretorio = input(" 📁  Diretório: ")
    while not diretorio:
        diretorio = input(" ❌  Error , Diretório não pode ser Nulo \n  📁 Diretório: ")
    while os.path.exists(diretorio):
        diretorio = input(" ❌  Error , Diretório já existe \n  📁 Diretório: ")
    diretorio = diretorio.strip()
    # contador , é igual a 0
    contador = 0
    # inicia no modulo 1 , caso não entenda , veja o  Moon_Load_Loading_Texts['criar_diretório']['módulo-1'] , conforme percorrer , incrementa o contador e o modulo
    modulo = 1
    # enquanto True , ou seja enquanto verdade , seja verdade , faça o looping
    while True:
        # zip() , é um interador de conjunto de arrays , me permitindo percorrer duas arrays ao mesmo tempo;
        # zip(array1 , array2)
        # para i , e em zip(Array1 , Array2)
        # i percorre a Array1 = i = variavel (poderia qualquer outra coisa)
        # t percorre a Array2 = t = variavel (poderia qualquer outra coisa)
        for i , t in zip(Moon_Load_Loading , Moon_Load_Loading_Texts['criar_diretório'][f'módulo-{modulo}']):
            # conta 1 milisegundo
            time.sleep(.1)
            # incrimenta o contador a cada loop
            contador = contador + 1
            # imprime a mensagem mostrando o dado de cada array
            print(f"\r {i} {t} | {contador}%             ", end="")
            # conta 1 milisegundo
            time.sleep(.1)
            # enquanto contador for igual a 100
            while contador == 100:
                # reinicia o contador igualando ele a 0 novamente
                contador = 0
                # cria o diretório
                if modulo == 1:
                    os.makedirs(diretorio)
                if modulo == 2:
                    os.makedirs(f"{diretorio}-Word")
                if modulo == 3:
                    os.makedirs(f"{diretorio}-PDF")
                # incrementa o modulo , de 1 para 2 , de 2 para 3 , de 3 para 4 , (os módulos da array 'criar_diretório' , vão apenas até o módulo 3 , mas paramos no 4 pois ele para antes de percorrer , então após o 3 incrementa e para (break) )
                modulo = modulo + 1
                # se for igual a 4 
                if modulo == 4:
                    # quebra o looping
                    break
        # se o modulo for igual a 4 
        if modulo == 4:
            modulo = 1
            contador = 0
            # quebra o looping 
            break
    print("\n")
    print("🔚")
    break_question_progress = input(f"o 📁  Diretório Criado {diretorio} \n é o Desejado ? \n 🟢  [S] para sim \n 🔴  [N] para não \n Resposta: ")
# cria uma função (def) com nome (validacoes) com o recebimento de um parametro (parametro)
def validacoes(parametro):
    # se parametro for igual a 'n'
    if parametro == 'n':
        # retorne uma mensagem de '_STOP'
        return '_STOP'
    # validators.url(parametro) , isso verifica a url (se um URL)
    # se é realmente uma url (validators.url(parametro)) , se realmente for , então usa o 'e' (and) verifica ao mesmo tempo se a URL já verificada (validators.url) contem "youtube" , 'e' (and) "youtu" , para verificar se realmente a URL é do YouTube , se for
    if validators.url(parametro) and "youtube" in parametro and "youtu" in parametro:
        # retorne True
        return True
    # se não
    else:
        # returna False
        return False
# Cria uma Array para Alocar todos os Links Informados pelo Usuário
_LINKS = []
# Imprime a Mensagem de STOP que é [n]
print("\n PARA PARAR BASTA A QUALQUER MOMENTO PRESSIONAR [N] \n")
# primeiro input de link para recebimento do Link
link = input("Insira o Link: 🔗  ")
# enquanto (while) não (not) mensagem (like) , ou seja a negação de uma mensagem é a ausência dela ou seja , enquanto a mensagem estiver vazia
while not link:
    # retorne o mesmo input com a mensagem de error
    link = input("❌  Error Link Não pode ser Vazio \n Insira o Link: 🔗  ")
# passando desse (while) é porque ela não é vazia
# igua-lá a variavel (link) a ela mesma com .strip() , .strip() = retira todos os espaços anteriores e posteriores , tipo 'Texto    ' e '    Texto' , .strip() se tornam 'Texto'
link = link.strip()
# enquanto a função retornar 'False' , ou seja URL inválido (não é uma URL)
while validacoes(link) == False:
    # retorna o mesmo input com a mensagem de error
    link = input("❌  Error Link Inválido \n Insira o Link: 🔗  ")
# se a validacao retornar True
if validacoes(link) == True:
    # adicione na Lista
    _LINKS.append(link)
# após o primeiro link ser inserido na Lista
# inicia-se um looping , para que podemos pegar inúmeros links
# enquanto a função (validacoes) , é diferente (!=) , de '_STOP' ou seja o link (input) , que o usuário digitou não é um [n] , continua o looping
while validacoes(link) != '_STOP':
    # mostra o input de novo
    link = input("Insira o Link: 🔗  ")
    # enquanto (while) não (not) resposta (input) , enquanto o link estiver vazio
    while not link:
        # reiprime o input com a mensagem de error
        link = input("❌  Error Link Não pode ser Vazio \n Insira o Link: 🔗  ")
    link = link.strip() # passa o .strip() , novamente para que não existam espaços anteriores e nem posteriores
    # enquanto funcao retornar False (a URL não é valida)
    while validacoes(link) == False:
        # reiprima o input com a mensagem de error
        link = input("❌  Error Link Inválido \n Insira o Link: 🔗  ")
    # se a funcao retornar '_STOP' é porque o usuário digitou 'n' que significa para parar de receber links
    if validacoes(link) == '_STOP':
        # então quebra o recebimento
        break
    # insere na lista os links 
    _LINKS.append(link)

# inicia um contador em 1
contador = 1
# incia um modulo em 1 
modulo = 1
# quebra uma linha \n
print("\n")
# mostra os links capturados
print(" 🌐  SEUS LINKS   🌐  ")
# quebra outra linha
print("\n")
# percorre a Array com um laço de repetição
# for (para) i (variavel) in (no/na) Array | variavel (i) para cada elemento da Array 
for i in _LINKS:
    # imprima o contador e o link (i)
    print(f"{contador}  🔗  {i}  🌐")
    # incrementa o contador
    contador = contador + 1

# quebra a Linha
print("\n")
# se o contador for (!=) diferente de 0
if contador != 0:
    # sete ele agora como 0 (reutilização de variavel)
    contador = 0
# se modulo for (!=) diferente de 1
if modulo != 1:
    # sete ele como 1
    modulo = 1
# enquanto True , ou seja , aqui estamos em um looping inifinito , quando declaramos , while True: , porque ? 
# enquanto True for Verdade , ele continuará o loop infinitamente a menos que alguma condição , o quebre (break) , então: 
# enquanto (while) True for verdade 
while True: 
    # percorra duas variaveis ao mesmo tempo com um laço de repetição , for
    # para (for) i (variavel) , t (variavel) no compactar(Array 1 , Array 2)
    # i para cada elemento da Array 1 
    # t para cada elemento da Array 2
    for i , t in zip(Moon_Load_Loading , Moon_Load_Loading_Texts['processamento_de_link'][f'módulo-{modulo}']):
        # conta um milisegundo
        time.sleep(.1)
        # incrementa o contador
        contador = contador + 1
        # mostra a mensagem com (i) , (t) e o contador
        print(f"\r {i}  {t}  | {contador}%                                   ", end="")
        # conta um milisegundo
        time.sleep(.1)
        # enquanto contador for igual a 100
        while contador == 100:
            # reinici-o , setando ele novamente como 0
            contador = 0
            # e incremente o modulo
            modulo = modulo + 1
            # se modulo for igual a 4
            if modulo == 4:
                # quebre o looping for
                break
    # se modulo for igual a 4
    if modulo == 4:
        # sete o modulo como 1
        modulo = 1
        # sete o contador como 0 
        contador = 0
        # quebre o looping (break) while
        break
# O mesmo processo abaixo é o mesmo de cima
if contador != 0:
    contador = 0
if modulo != 1:
    modulo = 1
while True:
    for i , t in zip(Moon_Load_Loading , Moon_Load_Loading_Texts['processamento_de_video'][f'módulo-{modulo}']):
        time.sleep(.1)
        contador = contador + 1
        print(f"\r {i}  {t}  |  {contador}%                                          ", end="")
        time.sleep(.1)
        while contador == 100:
            contador = 0
            modulo = modulo + 1
            if modulo == 5:
                break
    if modulo == 5:
        break
if contador != 0:
    contador = 0
if modulo != 1:
    modulo = 1
while True:
    for i , t in zip(Moon_Load_Loading , Moon_Load_Loading_Texts['processamento_de_video_not_audio'][f'módulo-{modulo}']):
        time.sleep(.1)
        contador = contador + 1
        print(f"\r {i}  {t}  |  {contador}%                                          ", end="")
        time.sleep(.1)
        while contador == 100:
            contador = 0
            modulo = modulo + 1
            if modulo == 5:
                break
    if modulo == 5:
        break
if contador != 0:
    contador = 0
if modulo != 1:
    modulo = 1 
while True:
    for i , t in zip(Moon_Load_Loading , Moon_Load_Loading_Texts['processamento_de_audio'][f'módulo-{modulo}']):
        time.sleep(.1)
        contador = contador + 1
        print(f"\r {i}  {t}  |  {contador}%                                          ", end="")
        time.sleep(.1)
        while contador == 100:
            contador = 0
            modulo = modulo + 1
            if modulo == 5:
                break
    if modulo == 5:
        break
print("\n")
# enumera a Array Lista com (enumerate)
for num,v in enumerate(_LINKS):
    # variavel (video) é igual (=) , instância (YouTube) , com dois parametros , url (v) , e o progresso do vídeo (on_progress_callback=on_progress)
    video = YouTube(v , on_progress_callback=on_progress)
    # variavel (titulo) é igual a variavel (video) '.' - para acessar as subcamadas dessa variavel , (title) pega o titulo do Vídeo acessado
    titulo = video.title
    # variavel (imagem) é igual a variavel (video) '.' - para acessar as subcamadas dessa variavel , (thumbnail_url) pega a URL da imagem do vídeo
    imagem = video.thumbnail_url
    # para (for) variavel (kw) no variavel (video) '.' - para acessar as subcamadas dessa variavel , (keyswords) pega uma array de palavras-chaves do vídeo
    # kw para cada palavra chave (keywords)
    for kw in video.keywords:
        # adiciona na Array _KEYSWORDS , tanto em nome normal
        _KEYSWORDS.append(kw)
        # adiciona na Array _KEYSWORDS , tanto em lower() = minusculo
        _KEYSWORDS.append(kw.lower())
        # adiciona na Array _KEYSWORDS , tanto em capitalize() = primeira letra maiscula
        _KEYSWORDS.append(kw.capitalize())
        # adiciona na Array _KEYSWORDS , tanto em upper() = tudo maiscula
        _KEYSWORDS.append(kw.upper())
    # aloca a variavel (video) nela mesma video = video.
    # .streams acessa as manipulaçoes de mídia
    # get_highest_resolution() = pega na melhor e mais alta resolução
    # output_path = direciona o caminho do salvamento
    video = video.streams.get_highest_resolution().download(output_path=diretorio)
    # mensagem de já está concluído
    print("✅  Download Concluído !!!")
    # variavel (audio) = instância (YouTube) , v (variavel) , progresso (on_progress_callback=on_progress)
    audio = YouTube(v , on_progress_callback=on_progress)
    # variavel (audio) = variavel (audio)
    # .streams acessa as manipulações de mídia
    # filter() = filtra
    # only_audio = apenas audio
    # .first() = primeiro
    # download() = download
    # filename = nome do documento
    # out_path = caminho do documento
    audio = audio.streams.filter(only_audio=True).first().download(filename=f'{num}.mp3',output_path=diretorio)
    # mensagem de concluído
    print("✅  Download Concluído !!!")
    # variavel (mp4) = instância (YouTube) , v (variavel) , progresso (on_progress_callback=on_progress)
    mp4 = YouTube(v , on_progress_callback=on_progress)
    # variavel (mp4) = variavel (mp4) 
    # .streams acessa as manipulações de mídia
    # filter() = filtrar 
    # .first() = primeiro
    # download() = download
    # filename = nome do arquivo
    # output_path = caminho do documento
    mp4 = mp4.streams.filter(only_video=True).first().download(filename=f'{num}.mp4' , output_path=diretorio)
    # mostre a mensagem de concluído
    print("✅  Download Concluído !!!")
    # variavel (jpg) = requests (biblioteca) .get() (pegue) .content (container)
    # variavel (jpg) aloca a requisiçao
    # requests (biblioteca vinda `import requests`) , essa biblioteca simplismente faz a requisição do youtube a imagem (pede que o youtube retorne a imagem em binario)
    # get() = pegue , o get() espera o link da imagem (thumbnail_url)
    # .content = retorna em container
    # requests = pede a thumbnail_url , e devolve um binario  
    jpg = requests.get(imagem).content
    # com (with) abrir (open) , caminho/documento.jpg (caminho/documento.jpg , não precisa existir , mas existirá quando compilar esse codigo) , "wb" = write binary (escreva em binarios) as = aplique um apelido (arquivo)
    with open(f"{diretorio}/{num}.jpg" , "wb") as arquivo:
        # arquivo = arquivo , .write = escreva
        arquivo.write(jpg)
        # ou seja escreva em binario , o binario devolvido pelo requests , formando assim a imagem no nosso diretorio
    # variavel (model) = whisper (biblioteca) . load_model = carregue o modelo , 'base' , ou seja carregue o modelo base
    model = whisper.load_model('base')
    # variavel (transcry) = modelo (model) , transcribe (transcreva) , caminho/documento.mp3
    transcry = model.transcribe(f"{diretorio}/{num}.mp3")
    # cria documento Word
    docx = Document()
    # variavel (heiding) = variavel (docx) ,  .add_heading('texto',level do titulo = 1 para maior , 2 para medio , 3 para pequeno) = adiciona titulo
    heiding = docx.add_heading(titulo,level=1)
    # for (para) i (variavel) in (no) range() (intervalo)
    # quando adicionamos o heading , caso: add_heading('titulo do meu documento' , level=1)
    # teremos: 
    # [0] = t [7] = ''  [14] = ''
    # [1] = i [8] = d   [15] = d
    # [2] = t [9] = o   [16] = o
    # [3] = u [10] = '' [17] = c
    # [4] = t [11] = m   .
    # [5] = l [12] = e   .
    # [6] = o [13] = u   .
    # assim precisamos percorrer todos os itens do seu titulo para torna-lós negritos e como faze-mos isso ? 
    # for (para) i (variavel) in (no) , range() (intervalo) , tipo: range(0 , 10) = contará de 0 a 10
    # nesse caso pegamos heiding (variavel) .runs (cada palavra desse heiding) , tudo isso dentro de um len (length = comprimento) , tendo len(heading.runs) = retornará o número de letras do seu texto
    for i in range(len(heiding.runs)):
        # i será o número , quando usamos 
        # heiding.runs[i] , seria como 
        # heiding.runs[0] , heiding.runs[1] ,... até o ultimo
        heiding.runs[i].bold # .bold torna o texto em negrito (cada palavra)
        heiding.runs[i].font.color.rgb = RGBColor(0, 0, 0) # .font (acessa a fonte do texto) .color (acessa a cor da font) . rgb (rgb da cor da fonte) = RGBColor(0,0,0) = significa cor Padrão
    # .alignment = alinhamento do seu heading 
    heiding.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # variavel (thumbnail) = variavel (docx) .add_picture , adiciona imagem , caminho/imagem.jpg , width = Inches(6) , converte o width (largura) de Inches(6) , 6 polegadas
    thumbnail = docx.add_picture(f"{diretorio}/{num}.jpg" , width=Inches(6))
    # variavel (texto) = transcry['text'] , pega o texto transcrito do vídeo no whisper , no modelo basíco acima
    texto = transcry['text']
    # cria um paragrafo vazio (transcription = docx.add_paragraph)
    transcription = docx.add_paragraph()
    # para cada palavra do texto , crie uma variavel (palavra)
    for palavra in texto.split():
        # run , seria uma palavra , add_paragraph() `sem parametros` , cria um paragrafo vazio , add_run() , adiciona palavras nesse paragrafo
        run = transcription.add_run(f"{palavra}  ")
        # se a palavra esta na array (_KEYSWORDS)
        if palavra in _KEYSWORDS:
            # passa a palavra para negrito
            run.bold = True # True para negrito
            run.italic = True # True para italico
            run.underline = True # True para sublinha
            run.font.highlight_color = WD_COLOR_INDEX.GREEN # .font (fonte) . highlight_color (cor do marcador)
    # variavel (transcription) , .paragraph_format (formato do paragrafo) , .space_after (espaço posterior) , Pt() (biblioteca) converte número em pontos (pt)
    transcription.paragraph_format.space_after = Pt(12)
    # variavel (transcription) , .paragraph_format (formato do paragrafo) , .space_before (espaço anterior) , Pt() (biblioteca) converte o número em pontos (Pt)
    transcription.paragraph_format.space_before = Pt(12)
    # variavel (transcription) , .alignment (alinhamento do paragrafo)
    transcription.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # variavel (transcription) , paragraph_format (formato do paragrafo) , .line_spacing (espaçamento entre linhas)
    transcription.paragraph_format.line_spacing = 1.5
    # salva o documento , no caminho/documento.docx
    docx.save(f"{diretorio}-Word/{num}.docx")
    
    # cria uma variavel com nome de extensão '.txt'
    arvore_de_elementos_pic_de_imagens = "_XMl_PIC.txt"
    # cria uma variavel com nome de extensão '.txt'
    arvore_de_elementos = "_XMl.txt"
    # verifica se elas existem
    if os.path.exists(arvore_de_elementos_pic_de_imagens) and os.path.exists(arvore_de_elementos):
        # se já existem 
        # remova-as
        os.remove(arvore_de_elementos_pic_de_imagens)
        os.remove(arvore_de_elementos)
    
    # variavel (leitor_docx) = instância (Document()) , caminho/documento.docx (irá ler o documento selecionado dentro do parametro da instância do Document)
    leitor_docx = Document(f"{diretorio}-Word/{num}.docx")
    
    # cria uma array para alocar o rIds , idenficação individual de pictures 'pic'
    _rIds = []
    
    # paragrafos (variavel) para cada paragrafo do documento (leitor_docx.paragraphs)
    for paragrafos in leitor_docx.paragraphs:
        # runs para cada palavra do paragrafo (paragrafos.runs)
        for runs in paragrafos.runs:
            # abrir o documento com "w" = writer , encoding="utf-8" , aplique o apelido _XML_arquivo
            with open(arvore_de_elementos , "w" , encoding="utf-8") as _XML_arquivo:
                # apelido 
                # .write() , escreva
                # runs._r.xml: 
                # palavra (runs)
                # _r = `_` significa interno , r elemento subjacente <w:r> do XML
                # xml = retorna a representação em XML
                _XML_arquivo.write(runs._r.xml)
            # se "pic" esta em runs._r.xml
            if "pic" in runs._r.xml:
                # abrir o documento com o "w" = write , encoding = "utf-8" , aplique o apelido arquivo
                with open(arvore_de_elementos_pic_de_imagens , "w" , encoding="utf-8") as arquivo: 
                    # apelido 
                    # .write() , escreva
                    arquivo.write(runs._r.xml)
                # variavel (arvore_de_elementos_xml) , igual (=) , xml '.' acessa a subcamada etree (e de elemento , tree arvore) '.' acessa a subcamada ElementTree (elemento de Arvore) , '.' subcamada ElementTree
                # .fromstring() =  transforma uma string em uma arvore (da String)
                arvore_de_elementos_xml = xml.etree.ElementTree.ElementTree(xml.etree.ElementTree.fromstring(runs._r.xml))
                # variavel (raiz_da_arvore_de_elementos) , igual (=) , variavel (arvore_de_elementos_xml) que agora é uma arvora inteira de elementos
                raiz_da_arvore_de_elementos = arvore_de_elementos_xml.getroot() # getroot = pegue a raiz , pega a raiz da arvora para percorrer os galhos da arvore
                # i (variavel) para cada elemento (frutos do galho `'{http://schemas.openxmlformats.org/drawingml/2006/main}blip'`) 
                # '{http://schemas.openxmlformats.org/drawingml/2006/main}blip' -> isso é um galho da Arvore , quer ver todos os galhos ? 
                # apenas faça
                # for i in raiz_da_arvore_de_elementos.iter() , com o iter() sem nenhum galho , ele trará todos os galhos
                for i in raiz_da_arvore_de_elementos.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                    i.items() #busca os itens do iter , mostra todos os frutos do galho selecionado
                    embed = i.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed') # pega o fruto desejado
                    # aloca na Array
                    _rIds.append(embed)
    # i (variavel) para cada elemento da array de frutos
    for i in _rIds:
        # j (variavel) para cada valor , .part (parte do documento que esta sendo manipulado) , .rels (relações que são tudo , tipo: hiperlinks , tabelas , imagens...) , .values() (busca todos as coisas ligadas a esse documento (rels) , em listas)
        for j in leitor_docx.part.rels.values():
            # j (variavel) , '.' acessa a subcamada da variavel que agora é um dos elementos (rels) trazidos por .values() e acessa o rId
            # se o rId for igual a i (variavel) dos frutos
            if j.rId == i:
                # verifica se tem 'media' no j (variavel) , '.' subcamada .target_ref (parte de destino)
                if "media" in j.target_ref:
                    # variavel (image_txt) , cria um documento com extensão .txt
                    image_txt = "image.txt"
                    # variavel (blob_binary) , que aloca a j (variavel) , '.' acessa a subcamada target_part (parte de destino) , '.' subcamada .blob (tras o documento mídia em binario)
                    blob_binary = j.target_part.blob
                    # abra o documento (image_txt) com "w" = write escreva , enconding="utf-8" , aplique o apelido (as) txt_image
                    with open(image_txt , "w" , encoding="utf-8") as txt_image:
                        # apelido . escreva (write)
                        # o binario
                        txt_image.write(str(blob_binary))
                    # adiciona o nome do elemento trazido pelo .values() do rels com uma extensão png
                    image = f'{j.rId}.png'
                    # abra documento (image) com "wb" = write binary , escreva binario , aplique o apelido imagem:
                    with open(image , "wb") as imagem:
                        # apelido  . escreva (write)
                        imagem.write(blob_binary) # cria a imagem no nosso diretorio

    # cria uma array associativa (dicionario)
    Informações_do_documento_Word = {}
    # variavel igual a None
    _XML = None
    # variavel igual a None
    _STYLE_PARAGRAPHS = None
    # variavel igual a uma string vazia
    _RUNS = ""
    # variavel igual a None
    _SPACING_LINE = None
    # variavel igual a None
    _SPACE_AFTER = None
    # variavel igual a None
    _SPACE_BEFORE = None
    # variavel com extensão txt
    _INFO_VIEW = "verificação_do_documento_em_texto.txt"
    # variavel com extensão txt
    _INFO_PREVIEW_ARC = "arquitetura_do_fluxo_de_codigo.txt"
    # variavel com string vazia
    _ALIGNMENT = ""
    # variavel com string vazia
    _FONT_SIZE = ""
    # variavel com string vazia
    _FONT_COLOR = ""
    # variavel de contador inicial em 0
    _CONT = 0
    # função (def) para verificar a cor dos highlights_color
    def _HIGHLIGHTS(param):
        if param == WD_COLOR_INDEX.GREEN:
            return 'green'
        elif param == WD_COLOR_INDEX.BLACK:
            return 'black'
        elif param == WD_COLOR_INDEX.BLUE:
            return 'blue'
        elif param == WD_COLOR_INDEX.BRIGHT_GREEN:
            return 'bright_green'
        elif param == WD_COLOR_INDEX.DARK_BLUE:
            return 'dark_blue'
        elif param == WD_COLOR_INDEX.DARK_RED:
            return 'dark_red'
        elif param == WD_COLOR_INDEX.DARK_YELLOW:
            return 'dark_yellow'
        elif param == WD_COLOR_INDEX.GRAY_25:
            return 'gray'
        elif param == WD_COLOR_INDEX.GRAY_50:
            return 'gray'
        elif param == WD_COLOR_INDEX.INHERITED:
            return 'inherited'
        elif param == WD_COLOR_INDEX.PINK:
            return 'pink'
        elif param == WD_COLOR_INDEX.RED:
            return 'red'
        elif param == WD_COLOR_INDEX.TEAL:
            return 'teal'
    # função (def) para verificar o alinhamento do paragrafo
    def _ALIGNMENT_ENCODING(param):
        if param == WD_ALIGN_PARAGRAPH.CENTER:
            return TA_CENTER
        elif param == WD_ALIGN_PARAGRAPH.RIGHT:
            return TA_RIGHT
        elif param == WD_ALIGN_PARAGRAPH.LEFT:
            return TA_LEFT
        elif param == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return TA_JUSTIFY 
        else:
            return TA_JUSTIFY
    # variavel (paragrafos) para cada paragrafos do documento (leitor_docx.paragraphs)
    for paragrafos in leitor_docx.paragraphs:
        # incrementa o contador
        _CONT = _CONT + 1
        # se o paragrafo tiver o estilo (style) de nome , 'Heading 1'
        if paragrafos.style.name == 'Heading 1':
            # adiciona o texto do paragrafo que tem o estilo do nome de 'Heading 1' na variavel de None
            _STYLE_PARAGRAPHS = f"<h1><b>{paragrafos.text}</b></h1>"
        # se o paragrafos tiver o formato do paragrafo (paragraph_format) tiver a espaçamento de linha (line_spacing)
        if paragrafos.paragraph_format.line_spacing:
            # adiciona o espaçamento da linha em uma variavel None
            _SPACING_LINE = paragrafos.paragraph_format.line_spacing
        # se o paragrafo tiver o formato do paragrafo (paragraph_format) tiver o espaço posterior (space_after)
        if paragrafos.paragraph_format.space_after:
            # adiciona o espaço posterior na variavel de None
            _SPACE_AFTER = paragrafos.paragraph_format.space_after
        # se o paragrafo tiver o formato do paragrafo (paragraph_format) tiver o espaço anterior (space_before)
        if paragrafos.paragraph_format.space_before:
            # adicione o espaço anterior na variavel de None
            _SPACE_BEFORE = paragrafos.paragraph_format.space_before
        # se o paragrafo tiver alinhamento (alignment)
        if paragrafos.alignment:
            # adicione o alinhamento na funçao e adicione o resultado que a função retornou (return) na variavel de None
            _ALIGNMENT = _ALIGNMENT_ENCODING(paragrafos.alignment)
        # runs para cada palavra do paragrafo (paragrafos.runs)
        for runs in paragrafos.runs:
            # se a palavra tem tamanho de fonte (font.size)
            if runs.font.size:
                # adicione o tamanho da fonte na variavel de None
                _FONT_SIZE = runs.font.size
            # se a palavra tem cor de fonte (font.color)
            if runs.font.color:
                # adiciona a cor da fonte da variavel None
                _FONT_COLOR = runs.font.color.rgb
            # se "pic" picture estiver na _r.xml
            if "pic" in runs._r.xml:
                # abra o documento de extensão txt com "w" = write escreva , enconding="utf-8" , aplique o apelido:
                with open(_INFO_PREVIEW_ARC , "w" , encoding="utf-8") as arquivo:
                    # apelido . escreva (write)
                    arquivo.write(str(runs._r.xml))
                # variavel (arvore) = conversão de metodos de mapeamento (xml.etree.ElementTree.ElementTree) e transformação de String em Arvores de Elementos (xml.etree.ElementTree.fromstring)
                arvore = xml.etree.ElementTree.ElementTree(xml.etree.ElementTree.fromstring(runs._r.xml))
                # variavel (raiz) = arvora de Elementos , '.' subcamada getroot (pegue a raiz da Arvore)
                raiz = arvore.getroot()
                # variavel de string vazia
                _EMBED = ""
                # variavel de string vazia
                _CX = ""
                # variavel de string vazia
                _CY = ""
                # i (variavel) para cada fruto do galho ('{http://schemas.openxmlformats.org/drawingml/2006/main}blip') da raiz (raiz = arvore.getroot) da arvore (arvore = xml.etree.ElementTree.ElementTree(xml.etree.ElementTree.fromstring(runs._r.xml)))
                for i in raiz.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                    # pega o fruto 
                    embed = i.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    # adiciona o fruto na variavel
                    _EMBED = embed
                # j (variavel) para cada fruto do galho ('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent') da raiz (raiz = arvore.getroot) da arvore (arvore = xml.etree.ElementTree.ElementTree(xml.etree.ElementTree.fromstring(runs._r.xml)))
                for j in raiz.iter('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'):
                    # pega o fruto e adiciona na variavel cy
                    cx = j.get('cx')
                    # pega o fruto e adiciona na variavel cx
                    cy = j.get('cy')
                    # adiciona a variavel na variavel de string vazia
                    _CX = cx
                    # adiciona a variavel na variavel de string vazia
                    _CY = cy
                # adiciona uma string concatenada do nome do fruto e seus cx e cy 
                _XML = f"{_EMBED}.{_CX}.{_CY}"
                # adiciona o a variavel de string vazia como string vazia , essa string vai concatenar cada palavra , se tem uma imagem não tem palavra então deixa vazia
                _RUNS = ""
            # se a palavra for negrita                        
            if runs.bold:
                # adiciona a variavel é igual a ela mesma mais concatenação da palavra em negrito
                _RUNS = _RUNS + f"<b>{runs.text}</b>"
            # se não , se a palavra for italico
            elif runs.italic:
                # adiciona a variavel é igual a ela mesma mais concatenação da palavra em italico 
                _RUNS = _RUNS + f"<i>{runs.text}</i>"
            # se não , se a palavra for sublinhada
            elif runs.underline:
                # adiciona a variavel é igual a ela mesma mais concatenação da palavra em sublinhada
                _RUNS = _RUNS + f"<u>{runs.text}</u>"
            # se não , se a palavra com marcador de cor
            elif runs.font.highlight_color:
                # adiciona a variavel é igual a ela mesma mais concatenação da palavra em o bgcolor (background color) da cor retornada pela função de verificação de cor de highlights_color
                _RUNS = _RUNS + f"<font bgcolor='{_HIGHLIGHTS(runs.font.highlight_color)}'>{runs.text}</font>"
            # se não , se a palavra tiver em negrito , e se a palavra tiver italico , e se tiver a palavra com marcação
            elif runs.bold and runs.italic and runs.underline and runs.font.highlight_color:
                # adiciona a variavel é igual a ela mesma mais concatenação da palavra em o bgcolor (background color) da cor retornada pela função de verificação de cor de hightlights_color , com negrito , com sublinhado e com italico
                _RUNS = _RUNS + f"<font bgcolor='{_HIGHLIGHTS(runs.font.highlight_color)}'><b><i><u>{runs.text}</u></i></b></font>"       
            # se não 
            else:
                # adiciona a variavel é igual a ela mesma mais concatenação da palavra
                _RUNS = _RUNS + f"<p>{runs.text}</p>"
            
            # Obs: porque eu "somo" a variavel com ela mesma mais a palavra , quando uma string vazia é somada com ela mesma e outra string , o valor anterior da string e concatenado com  o proximo valor formando de palavra em palavra uma frase e consequentemente um paragrafo
            
            # se a variavel definida para o espaçamento de linha , ainda continuar como none
            if _SPACING_LINE == None:
                # adicione o valor padrão 1.5 de espaçamento
                _SPACING_LINE = 1.5
            # se a variavel definida para o espaço posterior continuar como none
            if _SPACE_AFTER == None:
                # adiciona o valor padrão de 12
                _SPACE_AFTER = 12
            # se a variavel definida para o espaço anterior continuar como none
            if _SPACE_BEFORE == None:
                # adiciona o valor padrão de 12
                _SPACE_BEFORE = 12
            # se a variavel definida para o tamanho da fonte continuar como none
            if _FONT_SIZE == '':
                # adiciona o valor padrão de 12
                _FONT_SIZE = 12
            
            # adiciona na Array Associativa (dicionario)
            # adiciona a chave da array associativa como paragrafo - contador 
            Informações_do_documento_Word[f'paragrafo-{_CONT}'] = {
                # adiciona o paragrafo de Heading 1 que verificamos anteriormente
                'style_paragraph': _STYLE_PARAGRAPHS,
                # adiciona o espaçamento entre linhas que verificamos anteriormente
                'spacing_lines': _SPACING_LINE,
                # adiciona o espaço posterior que verificamos anteriormente
                'space_after': _SPACE_AFTER,
                # adiciona o espaço anterior que verificamos anteriormente
                'space_before': _SPACE_BEFORE,
                # adiciona o XML que verificamos anterioremente
                'XML': _XML,
                # adiciona o alinhamento que verificamos anteriormente
                'alignment': _ALIGNMENT,
                # adiciona o tamanho da fonte que verificamos anteriormente
                'font-size': _FONT_SIZE,
                # adiciona a cor da fonte que verificamos anteriormente
                'font-color': _FONT_COLOR,
                # adiciona o texto concatenado que concatenamos anteriormente
                'texto': _RUNS
            }
            # abra o documento com "w" = write , encoding="utf-8" , aplique o apelido (as) arquivo
            with open(_INFO_VIEW , "w" , encoding="utf-8") as arquivo:
                # apelido escreva informações da array
                arquivo.write(str(Informações_do_documento_Word))
    # cria uma array vazia
    informacoes = []
    # funções (def) para decodificar alinhamentos
    def _ALIGNMENT_DECODING(param):
        if param == 0:
            return WD_ALIGN_PARAGRAPH.LEFT
        elif param == 1:
            return WD_ALIGN_PARAGRAPH.CENTER
        elif param == 2:
            return WD_ALIGN_PARAGRAPH.RIGHT
        elif param == 3:
            return WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            return WD_ALIGN_PARAGRAPH.LEFT
    # cria uma variavel (pdf) com 
    # SimplesDocTemplade = Templade Documento Simples
    # de Nome = Caminho = Pasta/Nome_Do_Documento.pdf , pagesize=A4 , tamanho da folha será A4
    pdf = SimpleDocTemplate(f'{diretorio}-PDF/{num}.pdf',pagesize=A4)
    # key (variavel) , value (variavel) na Array 
    # .items() pega apenas o valores da array , como os proprios valores são chave-valor
    for key,value in Informações_do_documento_Word.items():
        # se o valor do style_paragraph não for Nulo
        if value['style_paragraph']:
            # cria uma variavel que alocará o estilo desse paragrafo atual
            # com ParagraphStyle()
            # name=Documento
            # leading = espaçamento entre linhas
            # fontSize = tamanho da Fonte
            # alingment = alinhamento
            # espaço anterior é o valor divido por 12700 por que é o valor de twips / pontos 
            # espaço posterior é o valor divido por 12700 por que é o valor de twips / pontos
            estilos_do_paragrafo = ParagraphStyle(
            name='DocumentoPDFCriado',
            leading= 17,
            fontSize= 16,
            alignment=_ALIGNMENT_DECODING(value['alignment']),
            spaceBefore= int(value['space_before']) / 12700,
            spaceAfter= int(value['space_after']) / 12700
        )
        # variavel (paragrafo_pdf) = a Paragraph(texto , estilo)
        # Paragraph(texto , estilo) = cria o paragrafo com o texto adicionado e o estilo
        paragrafo_pdf = Paragraph(value['style_paragraph'] , estilos_do_paragrafo)
        # adiciona na informação (array) o paragrafo
        informacoes.append(paragrafo_pdf)
        # se o valor do XML não for vazio
        if value['XML']:
            # variavel (png) = igual a uma string str() do valor com .split('.') que transforma uma string em array atraves de cada ponto na string e o indice [0] que pega o nome do fruto
            png = str(value['XML']).split('.')[0]
            # variavel (png) = igual a uma string str() do valor com .split('.') que transforma uma string em array atraves de cada ponto na string e o indice [1] que pega o valor do cx
            cx = str(value['XML']).split('.')[1]
            # variavel (png) = igual a uma string str() do valor com .split('.') que transforma uma string em array atraves de cada ponto na string e o indice [2] que pega o valor do cy
            cy = str(value['XML']).split('.')[2]
            # cria uma variavel (imagem) é igual (=) , Image(documento.png , width= , height=)
            imagem = Image(f"{png}.png" , width=400 , height=400)
            # adiciona na informacoes (array) a imagem
            informacoes.append(imagem)
            
        # cria uma variavel que alocará o estilo desse paragrafo atual
        # com ParagraphStyle()
        estilos_do_paragrafo = ParagraphStyle(
            # nome do paragrafo
            name='DocumentoPDFCriado',
            # leading é o espaçamento entre linhas que vai ser o valor do espaçamento entre linhas vezes o tamanho da fonte para ir adicionando uma atrás da outra
            leading= int(value['spacing_lines']) * int(value['font-size']),
            # alinhamento é o valor decodificado da função
            alignment=_ALIGNMENT_DECODING(value['alignment']),
            # espaço anterior é o valor divido por 12700 por que é o valor de twips / pontos 
            spaceBefore= int(value['space_before']) / 12700,
            # espaço posterior é o valor divido por 12700 por que é o valor de twips / pontos
            spaceAfter= int(value['space_after']) / 12700 
        )
         # variavel (paragrafo_pdf) = a Paragraph(texto , estilo)
        # Paragraph(texto , estilo) = cria o paragrafo com o texto adicionado e o estilo
        paragrafo_pdf = Paragraph(value['texto'] , estilos_do_paragrafo)
        # adiciona no informacoes (array) o paragrafo
        informacoes.append(paragrafo_pdf)
        
        # variavel (pdf) , .build(informacoes) , adiciona a array na build do documento (pdf) , para a criação do documento
        pdf.build(informacoes)
        
# https://www.youtube.com/watch?v=3J1xBL7zJXg&ab_channel=ProgramadorPython